# -*- coding: utf-8 -*-
"""
Created on Fri Jan 21 13:29:23 2022

@author: aso
"""

from sympy import Symbol
from docx import Document

from little_helpers import math_to_word


x = Symbol('x')
y = Symbol('y')

f_x = 3*x**2/(0.3*x**4-12)

# Write a single expression directly to a Word file
math_to_word(y, equals=f_x, mode='SymPy',
             word_file='math_to_word_example_1.docx',
             replace={'y': 'y_{c,r}', 'x': 'x_{org}'})

# Write multiple expressions to a single Word file
expr_1 = math_to_word(
    y, equals=f_x, mode='SymPy', replace={'y': 'y_{c,r}', 'x': 'x_{org}'})

expr_2 = math_to_word(
    y, equals=f_x-33+3*x/(4*y), mode='SymPy',
    replace={'y': 'y_{c,r}', 'x': 'x_{sub}^{sup}'})

doc = Document()
p = doc.add_paragraph()
p._element.append(expr_1)
p = doc.add_paragraph()
p._element.append(expr_2)
doc.save('math_to_word_example_2.docx')